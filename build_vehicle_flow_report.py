from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Vehicle_Mechanic_Assistance_System_Complete_Flow.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_header_row(table, labels):
    row = table.rows[0]
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_row(table, values):
    row = table.add_row()
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        p = cell.paragraphs[0]
        p.text = str(value)
        for run in p.runs:
            run.font.size = Pt(9)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(46, 116, 181)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(46, 116, 181)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.text = ""
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    r.font.size = Pt(10.5)
    p.add_run("\n" + body)
    for run in p.runs:
        run.font.name = "Calibri"
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10
for style_name in ["List Bullet", "List Number"]:
    styles[style_name].font.name = "Calibri"
    styles[style_name].font.size = Pt(11)
    styles[style_name].paragraph_format.space_after = Pt(8)
    styles[style_name].paragraph_format.line_spacing = 1.167

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Vehicle Mechanic Assistance System")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(11, 37, 69)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Complete Appointment and Emergency Flow")
r.font.name = "Calibri"
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(85, 85, 85)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Prepared from Flutter frontend code, WebSocket contracts, and attached backend snippets")
r.italic = True
r.font.name = "Calibri"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(85, 85, 85)

add_heading(doc, "What The App Does", 1)
add_body(
    doc,
    "This application is a vehicle mechanic assistance platform that connects users with nearby mechanics for two main service paths: scheduled appointments and emergency roadside assistance. Users can select their location, choose a service type such as bike mechanic, car mechanic, or puncture repair, submit problem details, track the mechanic, receive real-time status updates, approve or pay charges, and submit feedback. Mechanics receive requests through real-time notifications, accept or reject work, navigate to the customer, update job progress, send charges, confirm cash payment, and complete the job.",
)

add_callout(
    doc,
    "Source note",
    "The Flutter frontend project was available in the workspace. Backend behavior is described from the attached backend snippets and from the API endpoints, payloads, statuses, and WebSocket topics used by the frontend.",
)

add_heading(doc, "Main Actors And System Parts", 1)
for item in [
    "User: books appointments, creates emergency requests, tracks mechanics, approves charges, pays cash, and reviews service.",
    "Mechanic: receives appointment and emergency requests, accepts or rejects them, navigates to the user, updates job status, sends charges, and confirms payment.",
    "Backend API: stores users, mechanics, appointments, service requests, notifications, statuses, prices, payments, earnings, and reviews.",
    "WebSocket/STOMP layer: pushes real-time appointment, emergency, live-location, cancellation, payment, and completion events.",
    "Firebase/FCM layer: supports app notification setup, while the active workflows mostly rely on backend WebSocket topics.",
]:
    add_bullet(doc, item)

add_heading(doc, "Appointment Flow", 1)
add_heading(doc, "1. User Creates An Appointment", 2)
for item in [
    "The user opens the appointment booking screen and selects or picks a location on the map.",
    "The app sends the location to POST /api/user/bookappointment/nearbymechanics to fetch nearby mechanics.",
    "The user selects a service type: All, Puncher, Bike Mechanic, or Car Mechanic.",
    "The user selects appointment date and time, enters address and problem details.",
    "If the user chooses auto assignment, the app calls POST /api/user/auto/bookappointment.",
    "If the user manually selects a mechanic, the app calls POST /api/user/manual/bookappointment and sends the mechanic id.",
    "The backend creates an appointment, usually starting at PENDING, and returns appointment details or an appointment id.",
    "The app shows a booking confirmation screen with appointment id, service type, date, time, address, problem, and mechanic information if available.",
]:
    add_number(doc, item)

add_heading(doc, "2. Mechanic Receives And Accepts Appointment", 2)
for item in [
    "The mechanic dashboard initializes the mechanic WebSocket using the logged-in mechanic id.",
    "The mechanic subscribes to /topic/bookappointment/nearbymechanics/{mechanicId}.",
    "When a new appointment arrives, the mechanic gets an overlay or appointment notification.",
    "The mechanic opens Appointment Requests. The app fetches GET /api/mechanic/appointments/showmechanicappointments.",
    "PENDING appointments appear under Booking Requests with Accept and Reject actions.",
    "Accept calls GET /api/user/appointment/acceptappointment/{appointmentId}. Backend marks the appointment ACCEPTED and notifies the user.",
    "Reject calls GET /api/mechanic/appointment/rejectappointment/{appointmentId}. Backend marks the appointment REJECTED or informs the user that the request was rejected.",
]:
    add_number(doc, item)

add_heading(doc, "3. Mechanic Starts Travel", 2)
for item in [
    "After acceptance, the appointment appears in Upcoming.",
    "The mechanic taps Start Appointment.",
    "The frontend calls GET /api/mechanic/appointment/startappointment/{appointmentId}.",
    "The backend changes the status to ON_THE_WAY and sends a user notification on /topic/appointment/on-the-way/{userId}.",
    "The user appointment screen shows an ON THE WAY banner and can open tracking or contact the mechanic.",
]:
    add_number(doc, item)

add_heading(doc, "4. Mechanic Arrives", 2)
for item in [
    "When status is ON_THE_WAY, the mechanic card shows I Have Arrived.",
    "The frontend calls GET /api/mechanic/appointment/arrived/{appointmentId}.",
    "The attached backend code verifies the logged-in mechanic, finds the appointment for that mechanic, rejects invalid transitions, then sets AppointmentStatus.ARRIVED.",
    "The backend creates a MECHANIC_ARRIVED notification and sends /topic/appointment/arrived/{userId}.",
    "The user sees a Mechanic Arrived status message.",
    "The mechanic card should now show Start Work because the required next status is ARRIVED.",
]:
    add_number(doc, item)

add_heading(doc, "5. Work Starts And Completes", 2)
for item in [
    "The mechanic taps Start Work.",
    "The frontend calls GET /api/mechanic/appointment/startwork/{appointmentId}.",
    "The attached backend code only allows this if the current appointment status is ARRIVED, then sets IN_PROGRESS and notifies /topic/appointment/in-progress/{userId}.",
    "The mechanic taps Work Completed.",
    "The frontend calls GET /api/mechanic/appointment/completework/{appointmentId}.",
    "The attached backend code only allows this if the current status is IN_PROGRESS, then sets WORK_COMPLETED and notifies /topic/appointment/completework/{userId}.",
    "The user screen displays a Work Completed message and waits for the mechanic to send charges.",
]:
    add_number(doc, item)

add_heading(doc, "6. Mechanic Sends Charges And User Pays", 2)
for item in [
    "When status is WORK_COMPLETED, the mechanic card shows Send Charges.",
    "The mechanic enters the repair amount in a bottom-sheet popup.",
    "The frontend posts to /api/mechanic/appointment/sencharges/{appointmentId} with appid and finalPrice.",
    "The attached backend code sets repairAmount, amount = visitingCharge + repairAmount, paymentStatus = PENDING, and appointment status = PAYMENT_PROCESS.",
    "The backend sends /topic/appointment/sendcharges/{userId}.",
    "The user appointment card shows visiting charges, repair amount, total amount, and Pay Now.",
    "The user chooses Cash. The frontend calls GET /api/mechanic/appointment/paynow/{appointmentId}.",
    "The attached backend code sets paymentStatus = PAID, status = COMPLETED, paymentMethod = CASH, updates mechanic engagement, job count, and earnings, then sends /topic/appointment/appointmentdone/{mechanicId}.",
    "The mechanic receives Payment Done or Payment Received notification and the appointment moves to completed state.",
]:
    add_number(doc, item)

add_heading(doc, "7. Review After Appointment", 2)
for item in [
    "After cash payment succeeds, the user is navigated to ServiceReviewScreen.",
    "The review payload uses serviceType = APPOINTMENT.",
    "For appointments, the review sends appointmentId, rating, and comment.",
    "The endpoint is POST /api/service-request/review/submit.",
    "On success, the user returns to the home screen.",
]:
    add_number(doc, item)

add_heading(doc, "Appointment Status Timeline", 2)
status_table = doc.add_table(rows=1, cols=4)
add_header_row(status_table, ["Status", "Who Triggers It", "Frontend Action", "Backend Meaning"])
for row in [
    ("PENDING", "User books appointment", "Shows in mechanic Booking Requests", "Appointment is waiting for mechanic decision"),
    ("ACCEPTED", "Mechanic accepts", "Start Appointment becomes available", "Mechanic is assigned to appointment"),
    ("ON_THE_WAY", "Mechanic starts appointment", "User sees on-the-way banner; mechanic sees arrival button", "Mechanic is travelling to customer"),
    ("ARRIVED", "Mechanic taps arrived", "Mechanic sees Start Work", "Mechanic reached user location"),
    ("IN_PROGRESS", "Mechanic starts work", "Mechanic sees Work Completed", "Repair/service is currently being performed"),
    ("WORK_COMPLETED", "Mechanic completes work", "Mechanic sees Send Charges; user waits for bill", "Work is done but payment is not requested yet"),
    ("PAYMENT_PROCESS", "Mechanic sends charges", "User sees payment amount and Pay Now", "Payment is pending from user"),
    ("COMPLETED", "User pays cash", "Review screen opens; mechanic receives done event", "Job and payment are complete"),
    ("CANCELLED/REJECTED/EXPIRED", "User, mechanic, or backend timeout", "Moved to Cancelled tab", "Appointment no longer active"),
]:
    add_row(status_table, row)
set_table_width(status_table, [1.35, 1.45, 1.9, 1.8])

add_heading(doc, "Emergency Roadside Assistance Flow", 1)
add_heading(doc, "1. User Creates Emergency Request", 2)
for item in [
    "The user starts an emergency/service-request flow from the app.",
    "The user chooses service category such as bike, car, puncture, or other.",
    "The user can add notes and accept fixed inspection/visiting charges if required by the previous screen.",
    "The map screen resolves the selected location and builds a payload with serviceType, userNotes, isfixedchargeaccepted, user phone number, latitude, longitude, and location name.",
    "The app calls POST /api/service-request/create.",
    "The backend creates a service request and returns a requestId.",
    "The frontend stores the active request id and subscribes to /topic/request/{requestId}.",
    "The app fetches nearby mechanics using /api/service-request/nearbymechanic and shows the waiting state while mechanics are notified.",
]:
    add_number(doc, item)

add_heading(doc, "2. Mechanic Receives Emergency Request", 2)
for item in [
    "The mechanic WebSocket subscribes to /topic/mechanic/requests/{mechanicId} and also road request topics.",
    "If a payload contains userLatitude and requestId, the app opens MechanicRequestAlertScreen.",
    "The mechanic sees user information, service details, location, distance, timer, and accept action.",
    "Accept calls POST /api/service-request/accept/{requestId}.",
    "The backend assigns the request to that mechanic and returns merged request/mechanic/user data.",
    "The frontend saves the active tracking state and opens MechanicUserMap.",
    "Mechanic live location service starts and sends location updates to the backend.",
]:
    add_number(doc, item)

add_heading(doc, "3. Live Tracking And Arrival", 2)
for item in [
    "After acceptance, both sides listen to /topic/request/{requestId}.",
    "The mechanic map also listens to /topic/request/{requestId}/live-location for distance and ETA updates.",
    "The user map shows the accepted mechanic, route, distance, ETA, and mechanic movement.",
    "When the mechanic reaches the customer, the mechanic taps Have Arrived.",
    "The frontend sends GET /api/service-request/isarrived/{requestId} with current lat/lng query parameters.",
    "The backend verifies whether the mechanic is close enough to the customer. If not, the frontend shows distance such as You are X away from the customer.",
    "If arrival is accepted, the mechanic UI stops live tracking, marks hasArrived true, and allows charge entry.",
]:
    add_number(doc, item)

add_heading(doc, "4. Emergency Charges Approval", 2)
for item in [
    "After arriving, the mechanic opens Set Inspection Charges.",
    "The mechanic enters a final inspection/repair price.",
    "The frontend posts to /api/service-request/send-final-price with requestId and finalPrice.",
    "The backend sends a FINAL_PRICE_SENT event to the user request topic.",
    "The user screen receives the price and shows a review/confirm charges panel.",
    "The user approves the payment request by calling POST /api/service-request/approve-payment-request with requestId and finalPrice.",
    "After approval, the request status becomes APPROVED_PAYMENT_REQUEST or WORK_STARTED depending on backend naming.",
    "The mechanic side receives USER_APPROVED or approved status and can proceed with work.",
]:
    add_number(doc, item)

add_heading(doc, "5. Emergency Work Completion And Payment", 2)
for item in [
    "When work is finished, the mechanic calls GET /api/service-request/work-completed/{requestId}.",
    "The user screen receives WORK_COMPLETED or WAITING_FOR_PAYMENT and shows payment options.",
    "The user starts cash payment using POST /api/service-request/paynow/{requestId} with paymentype = CASH.",
    "The user screen enters PAYMENT_PENDING and asks the user to hand over cash.",
    "The mechanic confirms receiving cash using GET /api/service-request/completed/{requestId}.",
    "The backend completes the request, and the user receives PAYMENT_DONE or COMPLETED.",
    "The user is navigated to ServiceReviewScreen with serviceType = EMERGENCY.",
    "The mechanic returns to the dashboard after payment confirmation and job completion.",
]:
    add_number(doc, item)

add_heading(doc, "Emergency Status Timeline", 2)
em_table = doc.add_table(rows=1, cols=4)
add_header_row(em_table, ["State/Status", "User Side", "Mechanic Side", "Backend/Event Meaning"])
for row in [
    ("CREATED / WAITING", "User waits for mechanic", "Mechanics receive alert", "Request is open and broadcasted"),
    ("ACCEPTED", "Accepted mechanic appears on map", "Mechanic opens navigation map", "Request assigned to one mechanic"),
    ("ARRIVED", "User waits for charges/work", "Mechanic can send price", "Mechanic reached customer"),
    ("FINAL_PRICE_SENT", "User reviews price", "Mechanic waits for approval", "Mechanic submitted inspection or repair charge"),
    ("APPROVED_PAYMENT_REQUEST", "User approved charges", "Mechanic starts work", "Price approval received"),
    ("WORK_COMPLETED", "User sees payment due", "Mechanic waits for payment", "Work finished and payment is required"),
    ("PAYMENT_PENDING", "User has selected cash and should hand over money", "Mechanic confirms cash received", "Cash handover step is active"),
    ("COMPLETED / PAYMENT_DONE", "Review screen opens", "Mechanic returns to dashboard", "Emergency request is fully closed"),
    ("CANCELLED / EXPIRED", "User exits active flow", "Mechanic alert/map closes", "Request is no longer active"),
]:
    add_row(em_table, row)
set_table_width(em_table, [1.45, 1.75, 1.75, 1.55])

add_heading(doc, "Real-Time Notifications And Topics", 1)
topic_table = doc.add_table(rows=1, cols=3)
add_header_row(topic_table, ["Topic", "Receiver", "Purpose"])
for row in [
    ("/topic/bookappointment/nearbymechanics/{mechanicId}", "Mechanic", "New appointment request"),
    ("/topic/appointment/acceptappointment/{userId}", "User", "Appointment accepted"),
    ("/topic/appointment/final-reject/{userId}", "User", "Appointment finally rejected"),
    ("/topic/appointment/on-the-way/{userId}", "User", "Mechanic started travelling"),
    ("/topic/appointment/arrived/{userId}", "User", "Mechanic arrived"),
    ("/topic/appointment/in-progress/{userId}", "User", "Mechanic started work"),
    ("/topic/appointment/completework/{userId}", "User", "Work completed"),
    ("/topic/appointment/sendcharges/{userId}", "User", "Charges sent for appointment"),
    ("/topic/appointment/appointmentdone/{mechanicId}", "Mechanic", "Appointment payment completed"),
    ("/topic/mechanic/requests/{mechanicId}", "Mechanic", "Emergency service request alert"),
    ("/topic/request/{requestId}", "User and mechanic", "Emergency status, price, approval, cancellation, payment"),
    ("/topic/request/{requestId}/live-location", "User and mechanic", "Live distance and ETA updates"),
]:
    add_row(topic_table, row)
set_table_width(topic_table, [3.0, 1.25, 2.25])

add_heading(doc, "Important Backend Endpoints Used By The App", 1)
endpoint_table = doc.add_table(rows=1, cols=4)
add_header_row(endpoint_table, ["Area", "Method", "Endpoint", "Purpose"])
for row in [
    ("Appointment", "POST", "/api/user/bookappointment/nearbymechanics", "Find nearby mechanics for booking"),
    ("Appointment", "POST", "/api/user/auto/bookappointment", "Create appointment with auto assignment"),
    ("Appointment", "POST", "/api/user/manual/bookappointment", "Create appointment with selected mechanic"),
    ("Appointment", "GET", "/api/user/appointment/acceptappointment/{appointmentId}", "Mechanic accepts appointment"),
    ("Appointment", "GET", "/api/mechanic/appointment/rejectappointment/{appointmentId}", "Mechanic rejects appointment"),
    ("Appointment", "GET", "/api/mechanic/appointment/startappointment/{appointmentId}", "Move appointment to ON_THE_WAY"),
    ("Appointment", "GET", "/api/mechanic/appointment/arrived/{appointmentId}", "Move appointment to ARRIVED"),
    ("Appointment", "GET", "/api/mechanic/appointment/startwork/{appointmentId}", "Move appointment to IN_PROGRESS"),
    ("Appointment", "GET", "/api/mechanic/appointment/completework/{appointmentId}", "Move appointment to WORK_COMPLETED"),
    ("Appointment", "POST", "/api/mechanic/appointment/sencharges/{appointmentId}", "Send repair charges and move to PAYMENT_PROCESS"),
    ("Appointment", "GET", "/api/mechanic/appointment/paynow/{appointmentId}", "Cash payment and complete appointment"),
    ("Emergency", "POST", "/api/service-request/create", "Create emergency request"),
    ("Emergency", "POST", "/api/service-request/accept/{requestId}", "Mechanic accepts emergency request"),
    ("Emergency", "GET", "/api/service-request/isarrived/{requestId}", "Verify mechanic arrival with coordinates"),
    ("Emergency", "POST", "/api/service-request/send-final-price", "Mechanic sends final price"),
    ("Emergency", "POST", "/api/service-request/approve-payment-request", "User approves charges"),
    ("Emergency", "GET", "/api/service-request/work-completed/{requestId}", "Mechanic marks work completed"),
    ("Emergency", "POST", "/api/service-request/paynow/{requestId}", "User starts cash payment"),
    ("Emergency", "GET", "/api/service-request/completed/{requestId}", "Mechanic confirms cash and closes job"),
    ("Review", "POST", "/api/service-request/review/submit", "Submit appointment or emergency review"),
]:
    add_row(endpoint_table, row)
set_table_width(endpoint_table, [1.0, 0.7, 2.9, 1.9])

add_heading(doc, "Key Difference Between Appointment And Emergency", 1)
for item in [
    "Appointments are scheduled. They include date and time, can be auto-assigned or manually assigned, and progress through appointment-specific statuses from PENDING to COMPLETED.",
    "Emergency requests are immediate. They are map-centered, use live mechanic tracking, verify physical arrival using coordinates, and require user approval of charges before work continues.",
    "Appointment charge flow uses visitingCharge plus repairAmount. Emergency charge flow uses arrival/visiting price plus final inspection or service price.",
    "Appointment payment is completed from the user appointment history and notifies the mechanic on appointmentdone. Emergency payment is a cash handover flow where the user starts payment and mechanic confirms cash received.",
    "Both flows end with the same review screen, but appointment review sends serviceType = APPOINTMENT and appointmentId, while emergency review sends serviceType = EMERGENCY and serviceId.",
]:
    add_bullet(doc, item)

add_heading(doc, "End-To-End Summary", 1)
add_body(
    doc,
    "In the appointment journey, the user schedules a service, the mechanic accepts it, travels, arrives, starts and completes the work, sends charges, receives payment, and the user submits a review. In the emergency journey, the user creates an immediate roadside request, a mechanic accepts and is tracked live, arrival is verified by location, charges are sent and approved, work is completed, cash payment is confirmed, and the user submits a review. The whole system depends on synchronized backend statuses and WebSocket topics so both user and mechanic screens update in near real time.",
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Vehicle Mechanic Assistance System - Complete Flow")
fr.font.name = "Calibri"
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(85, 85, 85)

doc.save(OUT)
print(OUT)
